from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.5)

# ── Helper: set paragraph font ────────────────────────────────────────────────
def set_run(run, bold=False, size=11, color=None, italic=False):
    run.bold   = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = "Calibri"
    if color:
        run.font.color.rgb = RGBColor(*color)

def heading(doc, text, level=1, color=(26, 26, 26)):
    p = doc.add_heading(level=level)
    p.clear()
    run = p.add_run(text)
    set_run(run, bold=True, size=16 if level == 1 else (13 if level == 2 else 11),
            color=color)
    return p

def body(doc, text, space_after=6):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(space_after)
    for run in p.runs:
        run.font.name = "Calibri"
        run.font.size = Pt(11)
    return p

def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header row
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for run in hdr[i].paragraphs[0].runs:
            run.bold = True
            run.font.name = "Calibri"
            run.font.size = Pt(10)
        hdr[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        tc = hdr[i]._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "D4AF37")
        tcPr.append(shd)
    # Data rows
    for r_idx, row in enumerate(rows):
        cells = table.rows[r_idx + 1].cells
        for c_idx, val in enumerate(row):
            cells[c_idx].text = val
            for run in cells[c_idx].paragraphs[0].runs:
                run.font.name = "Calibri"
                run.font.size = Pt(10)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)
    return table

# ══════════════════════════════════════════════════════════════════════════════
#  COVER PAGE
# ══════════════════════════════════════════════════════════════════════════════
cover = doc.add_paragraph()
cover.paragraph_format.space_before = Pt(60)
cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = cover.add_run("MASHAAL CATERING SYSTEM")
set_run(r, bold=True, size=26, color=(180, 140, 30))

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = sub.add_run("Project Report")
set_run(r2, bold=False, size=16, color=(60, 60, 60))

doc.add_paragraph()
line = doc.add_paragraph("─" * 60)
line.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()
meta_lines = [
    ("System Type",  "Web-Based Catering Management System"),
    ("Architecture", "MVC (Model-View-Controller)"),
    ("Backend",      "PHP 8.x"),
    ("Database",     "MySQL / MariaDB"),
    ("Currency",     "Pakistani Rupee (PKR)"),
    ("Version",      "1.0.0"),
]
for label, val in meta_lines:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_label = p.add_run(f"{label}: ")
    set_run(r_label, bold=True, size=11, color=(100, 80, 10))
    r_val = p.add_run(val)
    set_run(r_val, size=11, color=(40, 40, 40))

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  1. INTRODUCTION
# ══════════════════════════════════════════════════════════════════════════════
heading(doc, "1. Introduction", 1, (180, 140, 30))
body(doc,
     "The Mashaal Catering System is a full-featured, web-based catering management platform "
     "developed using PHP 8 and MySQL. It is designed to streamline the operations of a catering "
     "business by providing two distinct portals — one for clients and one for staff/administrators. "
     "The system follows a strict Model-View-Controller (MVC) architectural pattern to ensure "
     "clean separation of concerns, maintainability, and scalability.")

body(doc,
     "Clients can register, browse menus, set up event details, build a custom catering order, "
     "and track their order history. Staff members can monitor incoming orders, update statuses, "
     "manage menus and packages, handle client records, and view business analytics — all from "
     "a single, intuitive dashboard.")

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
#  2. SYSTEM OVERVIEW
# ══════════════════════════════════════════════════════════════════════════════
heading(doc, "2. System Overview", 1, (180, 140, 30))

heading(doc, "2.1 Purpose", 2, (120, 95, 20))
body(doc,
     "The primary purpose of this system is to digitise the end-to-end workflow of a catering "
     "business — from client event setup and menu selection, through to order placement, payment "
     "tracking, and staff-side fulfillment management.")

heading(doc, "2.2 Scope", 2, (120, 95, 20))
body(doc, "The system covers the following functional areas:")
for item in [
    "Client self-registration and authentication",
    "Event setup and per-guest menu ordering",
    "Pre-set package selection and custom menu building",
    "Shopping cart with per-guest quantity scaling",
    "Order placement with payment method selection",
    "Staff order management and status updates",
    "Payment status tracking",
    "Menu and package management by staff",
    "Client directory management",
    "Revenue and analytics reporting",
]:
    p = doc.add_paragraph(item, style="List Bullet")
    for run in p.runs:
        run.font.name = "Calibri"
        run.font.size = Pt(11)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
#  3. USER ROLES
# ══════════════════════════════════════════════════════════════════════════════
heading(doc, "3. User Roles", 1, (180, 140, 30))
body(doc, "The system defines the following user roles:")
doc.add_paragraph()

add_table(doc,
    ["Role", "Description", "Entry Point"],
    [
        ["Client",  "End users who place catering orders",                "login_client.php"],
        ["Staff",   "Frontline employees who manage orders and payments",  "login_staff.php"],
        ["Manager", "Same access as Staff",                                "login_staff.php"],
        ["Admin",   "Full system access including client management",      "login_staff.php"],
    ],
    col_widths=[1.2, 3.5, 1.8]
)
doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
#  4. FEATURES
# ══════════════════════════════════════════════════════════════════════════════
heading(doc, "4. Features", 1, (180, 140, 30))

heading(doc, "4.1 Client Portal", 2, (120, 95, 20))
client_features = [
    ("User Registration & Login", "Clients can create an account and securely log in using bcrypt-hashed passwords."),
    ("Event Setup",               "Before browsing the menu, clients specify event type, date, time, location, guest count, and notes."),
    ("Pre-set Package Selection", "Clients can choose from curated catering packages that add all items to the cart at once."),
    ("Custom Menu Building",      "Clients can browse items by category and add individual items with per-guest quantity scaling."),
    ("Shopping Cart",             "A persistent session-based cart with item removal and running total displayed in PKR."),
    ("Checkout",                  "Clients select a payment method (Credit Card, Bank Transfer, Cash on Site) and confirm the order."),
    ("Order History",             "Clients can view all past orders with event details, amounts, and current statuses."),
    ("Profile Management",        "Clients can update their name, email, phone, and address at any time."),
]
for title, desc in client_features:
    p = doc.add_paragraph()
    r_title = p.add_run(f"{title}: ")
    set_run(r_title, bold=True, size=11)
    r_desc = p.add_run(desc)
    set_run(r_desc, size=11)
    p.paragraph_format.space_after = Pt(4)

doc.add_paragraph()
heading(doc, "4.2 Staff / Admin Portal", 2, (120, 95, 20))
staff_features = [
    ("Live Dashboard",        "Real-time stats showing pending orders, in-progress orders, finished orders, total revenue, upcoming events, and active clients."),
    ("Order Management",      "Staff can view all orders and update their status through a simple inline dropdown (Pending → Confirmed → In Progress → Finished → Cancelled)."),
    ("Payment Management",    "Payment statuses can be updated inline (Pending → Completed → Failed → Refunded)."),
    ("Menu & Package Mgmt",   "Staff can create and delete catering packages and add or remove individual menu items from packages."),
    ("Client Directory",      "A full client list with the ability to remove clients (soft or force-delete with full history wipe)."),
    ("Reports & Analytics",   "Monthly revenue bar chart, order status distribution, and top 5 most-ordered menu items with revenue breakdown."),
]
for title, desc in staff_features:
    p = doc.add_paragraph()
    r_title = p.add_run(f"{title}: ")
    set_run(r_title, bold=True, size=11)
    r_desc = p.add_run(desc)
    set_run(r_desc, size=11)
    p.paragraph_format.space_after = Pt(4)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
#  5. ARCHITECTURE
# ══════════════════════════════════════════════════════════════════════════════
heading(doc, "5. System Architecture", 1, (180, 140, 30))
body(doc,
     "The system is built around the MVC (Model-View-Controller) pattern. Each dashboard "
     "has a dedicated front controller at the root level that loads the appropriate controller "
     "(logic) and view (UI) files. Sub-pages are handled by action-specific controller files "
     "and rendered through include files within the view layer.")

heading(doc, "5.1 MVC Flow", 2, (120, 95, 20))
add_table(doc,
    ["Layer", "Location", "Responsibility"],
    [
        ["Front Controller", "admindash.php / clientdash.php",                "Entry point; loads controller + view"],
        ["Controller",       "controllers/admindashlogic.php",                "Session, POST handling, DB queries"],
        ["Controller",       "controllers/clientdashlogic.php",               "Session, POST handling, DB queries"],
        ["Sub-Controllers",  "controllers/staff_dashboard/*.php",             "Discrete action handlers (status, payment, etc.)"],
        ["Sub-Controllers",  "controllers/client_dashboard/events.php",       "Event setup, checkout, cart actions"],
        ["View",             "views/admindashUI.php",                         "Admin HTML/CSS template"],
        ["View",             "views/clientdashUI.php",                        "Client HTML/CSS template"],
        ["Sub-Views",        "views/staff_dashboard/*.php",                   "Menu management, reports views"],
        ["Sub-Views",        "views/client_dashboard/*.php",                  "Menu browser, cart/payment views"],
        ["Model",            "db.php",                                        "PDO connection + dbFetchOne/All/Execute helpers"],
    ],
    col_widths=[1.4, 2.8, 2.3]
)
doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
#  6. DATABASE SCHEMA
# ══════════════════════════════════════════════════════════════════════════════
heading(doc, "6. Database Schema", 1, (180, 140, 30))
body(doc, "The system uses the following relational database tables:")
doc.add_paragraph()

tables_info = [
    ("users",       ["id", "username", "password_hash", "full_name", "email", "role", "created_at"]),
    ("clients",     ["id", "full_name", "email", "password_hash", "phone", "address", "created_at"]),
    ("categories",  ["id", "name"]),
    ("menus",       ["id", "name", "description"]),
    ("menu_items",  ["id", "menu_id (FK)", "category_id (FK)", "name", "description", "price", "is_available"]),
    ("events",      ["id", "type", "date", "time", "location", "guest_count"]),
    ("orders",      ["id", "client_id (FK)", "user_id (FK)", "event_id (FK)", "status", "total_amount", "notes", "created_at"]),
    ("order_items", ["id", "order_id (FK)", "menu_item_id (FK)", "quantity"]),
    ("payments",    ["id", "order_id (FK)", "amount", "payment_date", "method", "status"]),
]

add_table(doc,
    ["Table", "Key Columns"],
    [(t, ", ".join(cols)) for t, cols in tables_info],
    col_widths=[1.5, 5.0]
)
doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
#  7. PREREQUISITES
# ══════════════════════════════════════════════════════════════════════════════
heading(doc, "7. Prerequisites", 1, (180, 140, 30))
add_table(doc,
    ["Requirement", "Version", "Notes"],
    [
        ["PHP",        "8.0 or higher", "PDO and PDO_MySQL extensions must be enabled"],
        ["MySQL",      "5.7 or higher", "MariaDB 10.4+ also supported"],
        ["Web Server", "Any",           "Apache, Nginx, or PHP built-in server"],
        ["Browser",    "Modern",        "Chrome, Firefox, or Edge recommended"],
    ],
    col_widths=[1.5, 1.5, 3.5]
)
doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
#  8. INSTALLATION & SETUP
# ══════════════════════════════════════════════════════════════════════════════
heading(doc, "8. Installation & Setup", 1, (180, 140, 30))

heading(doc, "8.1 Database Configuration", 2, (120, 95, 20))
body(doc, "Open db.php at the project root and set your database credentials:")
code_block = doc.add_paragraph()
r = code_block.add_run(
    "define('DB_HOST', 'localhost');\n"
    "define('DB_NAME', 'mashaal_catering');\n"
    "define('DB_USER', 'root');    // your MySQL username\n"
    "define('DB_PASS', '');        // your MySQL password"
)
r.font.name = "Courier New"
r.font.size = Pt(10)
code_block.paragraph_format.left_indent = Inches(0.5)

heading(doc, "8.2 Running with PHP Built-in Server", 2, (120, 95, 20))
for step, text in [
    ("Step 1", "Open a terminal and navigate to the project root folder."),
    ("Step 2", 'Run: php -S localhost:8000'),
    ("Step 3", "Open http://localhost:8000/login_client.php for the Client Portal."),
    ("Step 4", "Open http://localhost:8000/login_staff.php for the Staff Portal."),
]:
    p = doc.add_paragraph()
    r_step = p.add_run(f"{step}: ")
    set_run(r_step, bold=True, size=11)
    r_text = p.add_run(text)
    set_run(r_text, size=11)
    p.paragraph_format.space_after = Pt(4)

heading(doc, "8.3 Running with XAMPP / WAMP", 2, (120, 95, 20))
for step, text in [
    ("Step 1", "Copy the project folder to htdocs (XAMPP) or www (WAMP)."),
    ("Step 2", "Start Apache and MySQL from the control panel."),
    ("Step 3", "Visit http://localhost/catering-system/login_client.php"),
]:
    p = doc.add_paragraph()
    r_step = p.add_run(f"{step}: ")
    set_run(r_step, bold=True, size=11)
    r_text = p.add_run(text)
    set_run(r_text, size=11)
    p.paragraph_format.space_after = Pt(4)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
#  9. TECH STACK
# ══════════════════════════════════════════════════════════════════════════════
heading(doc, "9. Technology Stack", 1, (180, 140, 30))
add_table(doc,
    ["Layer", "Technology"],
    [
        ["Backend Language", "PHP 8.x"],
        ["Database",         "MySQL / MariaDB"],
        ["Frontend",         "HTML5, Vanilla CSS, JavaScript"],
        ["Typography",       "Google Fonts — Outfit"],
        ["Charts",           "Chart.js (CDN)"],
        ["Architecture",     "MVC (Model-View-Controller)"],
        ["Currency",         "Pakistani Rupee (PKR)"],
        ["Authentication",   "PHP Sessions + bcrypt password hashing"],
    ],
    col_widths=[2.0, 4.5]
)
doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
#  10. SECURITY CONSIDERATIONS
# ══════════════════════════════════════════════════════════════════════════════
heading(doc, "10. Security Considerations", 1, (180, 140, 30))
security_points = [
    ("Password Hashing",       "All passwords are hashed using PHP's password_hash() with the PASSWORD_DEFAULT algorithm (bcrypt)."),
    ("Session Management",     "Sessions are regenerated upon login using session_regenerate_id(true) to prevent session fixation attacks."),
    ("Role-Based Access",      "The requireRole() and requireClientLogin() guard functions are called at the top of every controller to enforce access control."),
    ("Input Sanitisation",     "All user-supplied data is trimmed and cast to appropriate types before use. HTML output uses htmlspecialchars() to prevent XSS."),
    ("SQL Injection Prevention","All database queries use PDO prepared statements with parameterised values — no raw SQL string interpolation of user input."),
]
for title, desc in security_points:
    p = doc.add_paragraph()
    r_title = p.add_run(f"{title}: ")
    set_run(r_title, bold=True, size=11)
    r_desc = p.add_run(desc)
    set_run(r_desc, size=11)
    p.paragraph_format.space_after = Pt(5)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
#  11. CONCLUSION
# ══════════════════════════════════════════════════════════════════════════════
heading(doc, "11. Conclusion", 1, (180, 140, 30))
body(doc,
     "The Mashaal Catering System provides a complete, production-ready solution for managing "
     "catering operations digitally. By adopting an MVC architecture, the codebase is organised, "
     "maintainable, and easy to extend. The dual-portal design — one for clients and one for staff "
     "— ensures that each user type has a focused, role-appropriate experience. "
     "The system is built entirely on open-source technologies (PHP, MySQL) and can be deployed "
     "on any standard web hosting environment with minimal configuration.")

body(doc,
     "Future enhancements could include email notifications upon order placement, an invoice PDF "
     "generator, an SMS reminder system for upcoming events, and a mobile-responsive redesign "
     "using a CSS framework.")

# ══════════════════════════════════════════════════════════════════════════════
#  SAVE
# ══════════════════════════════════════════════════════════════════════════════
out_dir = r"e:\1. SE\Projects\Catering System\project report"
os.makedirs(out_dir, exist_ok=True)
out_path = os.path.join(out_dir, "Mashaal_Catering_System_Report.docx")
doc.save(out_path)
print(f"Report saved to: {out_path}")
